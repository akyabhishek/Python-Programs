from docx import Document

def remove_hindi_characters(docx_file):
  """
  This function reads a DOCX file and removes all Hindi characters from it.

  Args:
    docx_file: The path to the DOCX file.

  Returns:
    A new DOCX file with Hindi characters removed.
  """
  doc = Document(docx_file)
  new_paragraphs = []
  for paragraph in doc.paragraphs:
    new_text = ""
    for char in paragraph.text:
      if not is_hindi_character(char):
        new_text += char
    new_paragraphs.append(new_text)

  new_doc = Document()
  for paragraph in new_paragraphs:
    new_doc.add_paragraph(paragraph)

  # Save the new document with a different name to avoid overwriting the original
  new_doc.save(f"{docx_file[:-5]}_no_hindi.docx")

def is_hindi_character(char):
  """
  This function checks if a character is a Hindi character based on its Unicode range.

  Args:
    char: The character to check.

  Returns:
    True if the character is Hindi, False otherwise.
  """
  hindi_characters_range = (0x0900, 0x097F)
  return ord(char) in range(hindi_characters_range[0], hindi_characters_range[1] + 1)

# Example usage
docx_file = "D:\Development\Python-Programs\docxfile\targeth.docx"
remove_hindi_characters(docx_file)
print("Hindi characters removed successfully!")
